from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.alert import Alert
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium import webdriver
from docx import Document
import datetime as dt
import pandas as pd
import numpy as np
import random
import time
import os
from io import StringIO,BytesIO
from num2words import num2words
from docx.shared import Pt,Inches,RGBColor,Cm
from docx.oxml import OxmlElement,parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ALIGN_VERTICAL
import matplotlib.pyplot as plt
from docx.table import Table
from docx.oxml.ns import nsdecls
import streamlit as st
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.chrome.options import Options
import locale
from datetime import timedelta

locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

# Caminho para o chromedriver

icon_path = "images/signature.png"

st.set_page_config(
    page_title="Gerar Termo Confissão de Dívida",
    page_icon=icon_path,
    layout="wide",
    initial_sidebar_state="expanded"
)

def app(RU,CURSO,entrada,PriVenc,parcelas,ValorParce,valor_negoc,juridico): 

    caminho_chromedriver = "chromedriver.exe"
    servico = ChromeService(executable_path=caminho_chromedriver)

    # Configurar as opções do Chrome para executar em modo headless
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--window-size=1920,1080")  
    chrome_options.add_argument("--disable-gpu")  # Necessário para evitar erros em alguns sistemas
    chrome_options.add_argument("--disable-dev-shm-usage")

    # Inicializar o navegador Chrome com as opções configuradas
    navegador = webdriver.Chrome(service=servico, options=chrome_options)
    # navegador = webdriver.Chrome(service=servico)

    alerta=Alert(navegador)

    def listaAnegociar():
        
        while len(navegador.find_elements(By.ID,'DataTables_Table_0')) < 1:
            time.sleep(1)

        table = navegador.find_element(By.ID,'DataTables_Table_0')
        table_html = table.get_attribute('outerHTML')
        dfLnog = pd.read_html(StringIO(table_html),header=0)[0]
        dfLnog['links'] = None  # cria uma nova coluna 'links' com valores None

        abalink = table.find_element(By.XPATH,'//*[@id="DataTables_Table_0"]/tbody')
        rows = abalink.find_elements(By.TAG_NAME, "tr") # get all the rows in the table

        for index, row in enumerate(rows):
            link_elements = row.find_elements(By.TAG_NAME, "a") # get all the links in the current row
            for link in link_elements:
                href = link.get_attribute('href')
                if href is not None and 'Cadastro' in href and 'Negociacoes' in href:
                    dfLnog.at[index, 'links'] = href   # adiciona o link à linha correspondente no DataFrame

        return dfLnog

    def entradaAcordo(entrada,PriVenc,Parcela,ValorParce):
        # PriVenc=pd.to_datetime(PriVenc.replace("/","-"),dayfirst=True)
        # Parcela=12
        priParce = [{"Número Parcela": 1, "Texto": "(entrada / primeiro vencimento)", 'Vencimento': PriVenc, 'Valor': entrada}]

        df = pd.DataFrame(data=priParce)
        # df.set_index("Número Parcela", inplace=True)

        if Parcela>1:
            outras=[]
            # PriVenc=pd.to_datetime(PriVenc.replace("/","-"),dayfirst=True)
            for i in range(int(Parcela)):
                dia =PriVenc.day
                mes=(PriVenc.month)+1+i
                ano=PriVenc.year
                if mes>12:
                    mes=mes-12
                    ano=ano+1
                    vencimentos=f"{ano}-{mes}-{dia}"
                    outras.append({"Número Parcela":i,"Texto":"", 'Vencimento':vencimentos,"Valor":ValorParce})
                else:
                    vencimentos=f"{ano}-{mes}-{dia}"
                    outras.append({"Número Parcela":i,"Texto":"", 'Vencimento':vencimentos,"Valor":ValorParce})
        outras=pd.DataFrame(data=outras)
        # outras.set_index("Número Parcela", inplace=True)
        dfEnt=pd.concat([df,outras])
        dfEnt['Vencimento'] = pd.to_datetime(dfEnt['Vencimento']).dt.strftime("%d/%m/%Y")
        dfEnt["Número Parcela"] = range(1,len(dfEnt)+1)
        dfEnt['Valor']=(dfEnt['Valor'].str.replace(",",".")).astype(float)
        dfEnt.loc['Total'] = round(dfEnt[['Valor']].sum(),2)
        dfEnt.loc['Total']=dfEnt.loc['Total'].apply(lambda x:"" if pd.isna(x) else x)
        dfEnt['VALOR DO ACORDO POR EXTENSO']=dfEnt['Valor'].apply(valor_por_extenso)
        dfEnt['Valor']=dfEnt['Valor'].apply(lambda x: f"R$ {x}").str.replace(".",",")
        dfEnt.loc['Total','Vencimento']='TOTAL'  
        dfEnt["Número Parcela"]=dfEnt["Número Parcela"].apply(lambda x:round(x) if not isinstance(x, str) else x)
        dfEnt=dfEnt.reset_index(drop=True)
        # dfEnt=dfEnt[['Número Parcela','ID', 'Situação','Vencimento','Valor','VALOR DO ACORDO POR EXTENSO']]
        dfEnt.loc[0,"Número Parcela"] = str(dfEnt["Número Parcela"][0])+ "-" + dfEnt["Texto"][0]
        dfEnt=dfEnt.drop('Texto',axis=1)
        return dfEnt

    def ParcelasAcordo():
        navegador.find_element(By.CLASS_NAME,'collapse-link').click()
        table = navegador.find_element(By.XPATH,'//*[@id="panel-titulos"]/table')
        table_html = table.get_attribute('outerHTML')
        table_html = StringIO(table_html)
        dfParce = pd.read_html(table_html,header=0)[0]

        dfparce=dfParce[['Número', 'Parcela', 'Vencimento', 'Dias de atraso',
        'Valor original', 'Multa', 'Juros', 'Valor Corrigido']]

        # Converta explicitamente para o tipo de dado desejado (por exemplo, object)
        dfparce.loc[:, 'Número'] = dfparce['Número'].astype('str')
        dfparce.loc[:, 'Dias de atraso'] = dfparce['Dias de atraso'].astype('str')

        dfparce.loc[:,'Valor original']=(dfparce['Valor original'].str.replace("R$ ","").str.replace(".","").str.replace(",",".")).astype(float)

        dfparce.loc[:,'Multa']=(dfparce['Multa'].str.replace("R$ ","").str.replace(".","").str.replace(",",".")).astype(float)

        dfparce.loc[:,'Juros']=(dfparce['Juros'].str.replace("R$ ","").str.replace(".","").str.replace(",",".")).astype(float)

        dfparce.loc[:,'Valor Corrigido']=(dfparce['Valor Corrigido'].str.replace("R$ ","").str.replace(".","").str.replace(",",".")).astype(float)

        dfparce = dfparce.copy()

        dfparce.loc['Total', ['Valor original', 'Multa', 'Juros', 'Valor Corrigido']] = dfparce[['Valor original', 'Multa', 'Juros', 'Valor Corrigido']].sum()

        dfparce.loc['Total']=dfparce.loc['Total'].apply(lambda x:"" if pd.isna(x) else round(x,2))

        dfparce.loc[:,'Valor original']=dfparce['Valor original'].apply(lambda x: f"R$ {x}").str.replace(".",",")

        dfparce.loc[:,'Multa']=dfparce['Multa'].apply(lambda x: f"R$ {x}").str.replace(".",",")

        dfparce.loc[:,'Juros']=dfparce['Juros'].apply(lambda x: f"R$ {x}").str.replace(".",",")

        dfparce.loc[:, 'Valor Corrigido'] = dfparce['Valor Corrigido'].apply(lambda x: f"R$ {x}").str.replace(".",",")

        dfparce.loc['Total','Dias de atraso']='TOTAL'

        dfparce=dfparce.rename(columns={'Número':'ID'})

        return dfparce

    # Função para substituir texto em um parágrafo
    def substituir_texto(paragrafo, alvo, substituto):
    
        for run in paragrafo.runs:
            if alvo in run.text:
                print(alvo, substituto)
                run.text = run.text.replace(alvo, substituto)
            elif alvo in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(alvo, substituto)

    def valor_por_extenso(valor):
        try:
            parte_inteira, parte_decimal = str(valor).split('.')
        except:
            parte_inteira, parte_decimal = str(valor).split(',')

        # Converter a parte inteira para texto por extenso
        texto_inteiro = num2words(int(parte_inteira), lang='pt_BR').replace('-', 'negativo ')

        # Se houver parte decimal, converter também
        texto_decimal = ""
        if parte_decimal:
            try:
                texto_decimal = f" e {num2words(int(parte_decimal), lang='pt_BR').replace('-', 'negativo ')} centavos"
            except:
                texto_decimal = f" e {num2words(int(parte_decimal.replace("R$ ","")), lang='pt_BR').replace('-', 'negativo ')} centavos"
        # Adicionar a moeda
        moeda = "real" if parte_inteira == "1" else "reais"

        # Combinar as partes
        texto_completo = f"{texto_inteiro} {moeda}{texto_decimal}"

        return texto_completo

    def inserirTab(doc,tabela,base):
        # Encontrar o índice do parágrafo que contém o texto '{TABELA1}'
        indice_paragrafo_tabela = None
        for i, paragrafo in enumerate(doc.paragraphs):
            if tabela in paragrafo.text:
                indice_paragrafo_tabela = i
                break
        # Substituir o parágrafo pelo conteúdo da tabela
        if indice_paragrafo_tabela is not None:
            # Adicionar quebra de página antes da tabela
            paragrafo = doc.paragraphs[indice_paragrafo_tabela]
            paragrafo.clear()  # Limpar o parágrafo existente
            paragrafo.add_run().add_break()

            # Adicionar a tabela no índice desejado
            table = doc.add_table(rows=len(base) + 1, cols=len(base.columns))

            # Adicionar cabeçalho
            for j, coluna in enumerate(base.columns):
                cell = table.cell(0, j)
                cell.text = coluna
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.size = Pt(10)  # Ajuste o tamanho da fonte conforme necessário
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

            # Preencher dados
            for i, row in enumerate(base.itertuples(), start=1):
                for j, valor in enumerate(row[1:], start=0):  # Ajuste o índice inicial para 0
                    table.cell(i, j).text = str(valor)

            # Mover a tabela para o local correto
            tbl, p = table._tbl, paragrafo._p
            p.addnext(tbl)

            # Adicionar bordas
            border_xml = """
            <w:tcBorders %s>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tcBorders>
            """ % nsdecls('w')

            for row in table.rows:
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))

            # Definir a largura total da tabela
            table_width = Cm(16)  # Ajuste a largura total conforme necessário

            # Definir a largura das células
            num_cols = len(base.columns)
            cell_width = table_width / num_cols

            for row in table.rows:
                for cell in row.cells:
                    cell.width = cell_width

    def TipoDoc(juridico):
        if juridico=="Jurídico":
            arquivo_original = 'MODELO - Termo de Acordo ANTES DA SENTENÇA.docx'
            return arquivo_original
        elif juridico=="Padrão":  
            arquivo_original = 'MODELO - Termo de Acordo ANTES DA SENTENÇA.docx'
            return arquivo_original

    # Caminho para o arquivo original e novo arquivo
    arquivo_original=TipoDoc(juridico)
    novo_arquivo = f'TERMOS_EDITADOS_{random.randint(1, 1000)}.docx'

    # Abrir o documento existente
    doc = Document(arquivo_original)

    # Dados das variáveis dinâmicas

    navegador.get(f'https://sicob.uninter.com/Login')  
    navegador.get(f'https://sicob.uninter.com/Login')  
    navegador.find_element(By.XPATH,'/html/body/div/div/form/div/a/button').click()
    
    navegador.maximize_window()
    while len(navegador.find_elements(By.CLASS_NAME,'form-control-top-pesquisa')) < 1:
        time.sleep(0.5)
    navegador.find_element(By.CLASS_NAME,'form-control-top-pesquisa').send_keys(RU)
    time.sleep(0.5)
    navegador.find_element(By.CLASS_NAME,'form-control-top-pesquisa').send_keys(Keys.ENTER)

    dadossecob=navegador.find_elements(By.CLASS_NAME,"control-label-left")

    listaDados=[item.text for item in dadossecob][2:]

    dicionario_de_dados = {listaDados[i]: listaDados[i + 1] for i in range(0, len(listaDados), 2)}
    dfDados=pd.DataFrame(data=[item for item in dicionario_de_dados.items()],columns=['Tipo',"Informação"])
    nome_aluno = dicionario_de_dados['Nome']
    cpf = dicionario_de_dados['CPF']
    ru = dicionario_de_dados['RU']
    try:
        cel = dicionario_de_dados['Telefone Celular original']
    except:
        cel = dicionario_de_dados['Telefone Celular']
    email = dicionario_de_dados['E-mail original']
    endereco = f'{dicionario_de_dados['Endereço']} - {dicionario_de_dados['Bairro']}, {dicionario_de_dados['Cidade/UF']}, {dicionario_de_dados['CEP']}'

    dfLnog=listaAnegociar()

    time.sleep(2)
    dfLnog=dfLnog.loc[dfLnog['Turma'].str.contains(CURSO)]
    for linha2 in dfLnog.index:
        link=dfLnog.loc[linha2,'links']
        navegador.get(link)
        while len(navegador.find_elements(By.CLASS_NAME,'headings')) < 1:
            time.sleep(1)
        dfparce=ParcelasAcordo()
        
    dfEnt=entradaAcordo(entrada,PriVenc,parcelas,ValorParce)

    turma_curso = CURSO

    while len(navegador.find_elements(By.ID,'total_divida_atualizada')) < 1:
            time.sleep(1)

    valor_real = dfparce.loc['Total','Valor Corrigido']
    
    valor_extenso_real = valor_por_extenso(float(valor_real.split(" ")[1].replace(",",".")))
    print(valor_negoc)

    valor_extenso_negoc=valor_por_extenso(valor_negoc.split(" ")[0])
    print(valor_extenso_negoc)

    # Dicionário de substituição
    substituicoes = {
        "{TURMA_CURSO}": turma_curso,
        "{NOME_ALUNO}": nome_aluno,
        "{CPF}": cpf,
        "{RU}": ru,
        "{CEL}": cel,
        "{E-MAIL}": email,
        "{ENDERECO}": endereco,
        "{VALOR_REAL}": str(valor_real),
        "{VALOR_EXTENSO_REAL}": valor_extenso_real,
        "{VALOR_NEGOCIACAO}":str(valor_negoc),
        "{VALOR_EXTENSO_NEGOCIACAO}":valor_extenso_negoc,
        "{PARCELAS}": str(parcelas),
        "{VENCIMENTO}": PriVenc.strftime("%d/%m/%Y"),
        "{DATA_EXTENSO}": f"{str(dt.datetime.now().strftime("%d de %B de %Y"))}"
    }

    for i, paragrafo in enumerate(doc.paragraphs):
        for alvo, substituto in substituicoes.items():
            substituir_texto(paragrafo, alvo, substituto)

        # Considerar também o texto na caixa de texto do parágrafo
        if hasattr(paragrafo, 'text_frame'):
            for caixa_texto in paragrafo.text_frame.paragraphs:
                for alvo, substituto in substituicoes.items():
                    substituir_texto(caixa_texto, alvo, substituto)

    ## Salvar o novo documento
    doc.save(novo_arquivo)
    tabela1="{TABELA1}"
    tabela2="{TABELA2}"

    # Abrir o novo documento
    doc = Document(novo_arquivo)

    inserirTab(doc,tabela1,dfparce)

    inserirTab(doc,tabela2,dfEnt)

    doc.save(novo_arquivo)

    caminho_arquivo=os.getcwd()

    # Caminho completo para o arquivo
    caminho_arquivo = os.path.join(caminho_arquivo, novo_arquivo)

    # Salvar o documento
    doc.save(caminho_arquivo)

    st.success('Processo finalizado!', icon="✅")

    with st.spinner("Gerando Documento..."):
        st.download_button(
            label="Clique para baixar",
            data=open(caminho_arquivo, "rb").read(),
            file_name=f"Termo - {RU}-{nome_aluno}.docx",
            key="download_button",
        )
    return doc,dfparce,dfEnt,dfDados

# Abrir o novo documento no VSCode
# os.system(f'code {novo_arquivo}')
    
# def inputs():
#     RU = st.text_input('RU')
#     NEGOCIACAO=st.text_input('Nº Negociação')
#     st.button("Gerar Doc", type="primary")
#     if st.button('Say hello'):
#         app(RU,NEGOCIACAO)
#     return 

# RU=1602629
# NEGOCIACAO='3447725'
# turma_curso='adm'
# RU=4178402
# CURSO='2022/07 GD PSICANÁLISE'
# valor_negoc="820,50"
# entrada="100,00"
# PriVenc="2024-02-09"
# PriVenc=pd.to_datetime(PriVenc)
# parcelas=10
# ValorParce="150,00"

if __name__ == "__main__":
    col1,col2=st.columns([2,20])
    with col1:
        st.image("marca-uninter-horizontal.png", use_column_width=True)
    with col2:
        st.header("*Gerar termo de Confissão de Dívida*",divider='rainbow')
    col1,col2=st.columns([2,6])
    with col1:
        juridico = st.radio(
        "Selecione o Tipo do Documento",
        ["Jurídico", "Padrão"],
        index=None,
    )
        RU = st.text_input('RU')
        CURSO=st.text_input('Turma/Curso')
        valor_negoc=st.text_input('Valor Total Negociado')
        if len(valor_negoc.split(","))<2:
            valor_negoc=f"{valor_negoc},00"
        entrada=st.text_input('Valor Entrada')
        if len(entrada.split(","))<2:
            entrada=f"{entrada},00"
        PriVenc=st.date_input("Primeiro Vencimento",format="DD/MM/YYYY")
        parcelas = st.number_input('Nº Parcelas', format="%i", value=0)
        ValorParce=st.text_input('Valor Parcelas')
        if len(ValorParce.split(","))<2:
            ValorParce=f"{ValorParce},00"
        # st.button("Gerar Doc", type="primary")

        if st.button('Gerar') and RU !="":
            progress_text = "Operação em progresso. Por favor aguarde."
            my_bar = st.progress(0, text=progress_text)
            for percent_complete in range(100):
                time.sleep(0.01)
                my_bar.progress(percent_complete + 1, text=progress_text)
            time.sleep(1)
            my_bar.empty()
            doc,dfparce,dfEnt,dfDados=app(RU,CURSO,entrada,PriVenc,parcelas,ValorParce,valor_negoc,juridico)
            
        else:
            pass
    with col2:
        try:
            col1,col2=st.columns([22,14])
            with col1:
                st.dataframe(dfparce,hide_index=True)
            with col1:   
                st.dataframe(dfEnt,hide_index=True)
            with col2:
                st.dataframe(dfDados,hide_index=True)
                
        except:
            ""
#     st.download_button(
#     label="Download doc",
#     data=doc,
#     file_name='TERMOS_EDITADOS2.docx',
#     mime='text/doctx',
# )
    
